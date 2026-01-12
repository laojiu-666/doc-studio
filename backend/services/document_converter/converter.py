"""
Document converter service for converting between docx and HTML.
"""
import os
import uuid
import tempfile
from pathlib import Path

import mammoth
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .style_parser import StyleParser


class DocumentConverter:
    """Service for converting documents between formats."""

    def __init__(self):
        self.style_parser = StyleParser()

    def docx_to_html(self, file_path: str) -> str:
        """
        Convert a docx file to HTML for Tiptap editor.

        Args:
            file_path: Path to the docx file

        Returns:
            HTML string
        """
        with open(file_path, 'rb') as docx_file:
            result = mammoth.convert_to_html(docx_file)
            return result.value

    def html_to_docx(self, html_content: str, title: str) -> str:
        """
        Convert HTML content to a docx file.

        Args:
            html_content: HTML string from Tiptap editor
            title: Document title

        Returns:
            Path to the generated docx file
        """
        from bs4 import BeautifulSoup

        # Create a new document
        doc = Document()

        # Parse HTML
        soup = BeautifulSoup(html_content, 'html.parser')

        # Process page setup if present
        page_setup_elem = soup.find('page-setup')
        if page_setup_elem:
            attrs = {k: v for k, v in page_setup_elem.attrs.items()}
            page_setup = self.style_parser.parse_page_setup(attrs)
            self.style_parser.apply_page_setup(doc.sections[0], page_setup)

        # Process header if present
        header_elem = soup.find('header')
        if header_elem:
            section = doc.sections[0]
            header = section.header
            header.paragraphs[0].text = header_elem.get_text()

        # Process footer if present
        footer_elem = soup.find('footer')
        if footer_elem:
            section = doc.sections[0]
            footer = section.footer
            footer.paragraphs[0].text = footer_elem.get_text()

        # Process each element
        for element in soup.children:
            if element.name not in ['page-setup', 'header', 'footer']:
                self._process_element(doc, element)

        # Save to temp file
        temp_dir = tempfile.gettempdir()
        output_path = os.path.join(temp_dir, f'{uuid.uuid4()}.docx')
        doc.save(output_path)

        return output_path

    def _process_element(self, doc: Document, element):
        """Process an HTML element and add it to the document."""
        if element.name is None:
            # Text node
            text = str(element).strip()
            if text:
                doc.add_paragraph(text)
            return

        # Get element styles
        styles = {}
        if hasattr(element, 'get') and element.get('style'):
            styles = self.style_parser.parse_style_string(element.get('style'))

        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            heading = doc.add_heading('', level=level)
            self._process_inline(heading, element)
            if styles:
                para_style = self.style_parser.parse_paragraph_style(styles)
                self.style_parser.apply_paragraph_style(heading, para_style)

        elif element.name == 'p':
            para = doc.add_paragraph()
            self._process_inline(para, element)
            if styles:
                para_style = self.style_parser.parse_paragraph_style(styles)
                self.style_parser.apply_paragraph_style(para, para_style)

        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                para = doc.add_paragraph(style='List Bullet')
                self._process_inline(para, li)
                if li.get('style'):
                    li_styles = self.style_parser.parse_style_string(li.get('style'))
                    para_style = self.style_parser.parse_paragraph_style(li_styles)
                    self.style_parser.apply_paragraph_style(para, para_style)

        elif element.name == 'ol':
            for li in element.find_all('li', recursive=False):
                para = doc.add_paragraph(style='List Number')
                self._process_inline(para, li)
                if li.get('style'):
                    li_styles = self.style_parser.parse_style_string(li.get('style'))
                    para_style = self.style_parser.parse_paragraph_style(li_styles)
                    self.style_parser.apply_paragraph_style(para, para_style)

        elif element.name == 'blockquote':
            para = doc.add_paragraph()
            para.style = 'Quote'
            self._process_inline(para, element)
            if styles:
                para_style = self.style_parser.parse_paragraph_style(styles)
                self.style_parser.apply_paragraph_style(para, para_style)

        elif element.name == 'table':
            self._process_table(doc, element)

        elif element.name in ['div', 'section', 'article']:
            # Container elements - process children
            for child in element.children:
                self._process_element(doc, child)

    def _process_inline(self, paragraph, element):
        """Process inline elements within a paragraph."""
        for child in element.children:
            if child.name is None:
                # Text node
                text = str(child)
                if text:
                    paragraph.add_run(text)
            elif child.name == 'strong' or child.name == 'b':
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name == 'em' or child.name == 'i':
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif child.name == 'u':
                run = paragraph.add_run(child.get_text())
                run.underline = True
            elif child.name == 's' or child.name == 'strike':
                run = paragraph.add_run(child.get_text())
                run.font.strike = True
            elif child.name == 'span':
                run = paragraph.add_run(child.get_text())
                if child.get('style'):
                    span_styles = self.style_parser.parse_style_string(child.get('style'))
                    font_style = self.style_parser.parse_font_style(span_styles)
                    self.style_parser.apply_font_style(run, font_style)
            elif child.name == 'a':
                # Links - add text with underline
                run = paragraph.add_run(child.get_text())
                run.underline = True
            elif child.name == 'br':
                paragraph.add_run('\n')
            else:
                # Recursively process nested inline elements
                self._process_inline(paragraph, child)

    def _process_table(self, doc: Document, table_element):
        """Process an HTML table element."""
        rows = table_element.find_all('tr')
        if not rows:
            return

        # Count columns from first row
        first_row = rows[0]
        cols = len(first_row.find_all(['td', 'th']))

        if cols == 0:
            return

        # Create table
        table = doc.add_table(rows=len(rows), cols=cols)
        table.style = 'Table Grid'

        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for j, cell in enumerate(cells):
                if j < cols:
                    table.rows[i].cells[j].text = cell.get_text()
