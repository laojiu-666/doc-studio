"""
Incremental document converter for efficient docx updates.
"""
import os
import uuid
import hashlib
import tempfile
import copy
from dataclasses import dataclass, field
from typing import List, Optional, Dict, Any, Tuple
from enum import Enum

from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.shared import Pt, Inches

from .style_parser import StyleParser, FontStyle, ParagraphStyle, PageSetup


class ElementType(Enum):
    """Document element types."""
    HEADING = 'heading'
    PARAGRAPH = 'paragraph'
    LIST_ITEM = 'list_item'
    TABLE = 'table'
    BLOCKQUOTE = 'blockquote'
    PAGE_SETUP = 'page_setup'
    HEADER = 'header'
    FOOTER = 'footer'


class ChangeType(Enum):
    """Types of changes between document versions."""
    INSERT = 'insert'
    DELETE = 'delete'
    MODIFY = 'modify'
    MOVE = 'move'


@dataclass
class DocumentElement:
    """Represents a document element (paragraph, heading, etc.)."""
    type: ElementType
    content: str
    html: str
    level: int = 0  # For headings (1-6) or list nesting
    list_type: str = ''  # 'bullet' or 'number'
    styles: Dict[str, str] = field(default_factory=dict)
    children: List['DocumentElement'] = field(default_factory=list)  # For tables

    def get_hash(self) -> str:
        """Get content hash for comparison."""
        content = f"{self.type.value}:{self.level}:{self.content}:{self.list_type}"
        return hashlib.md5(content.encode()).hexdigest()[:16]


@dataclass
class Change:
    """Represents a change between document versions."""
    type: ChangeType
    index: int  # Position in document
    old_element: Optional[DocumentElement] = None
    new_element: Optional[DocumentElement] = None


class DocumentStructure:
    """Parse HTML into structured document representation."""

    def __init__(self, html: str):
        self.html = html
        self.elements: List[DocumentElement] = []
        self.page_setup: Optional[PageSetup] = None
        self.header_content: str = ''
        self.footer_content: str = ''
        self._parse()

    def _parse(self):
        """Parse HTML into elements."""
        soup = BeautifulSoup(self.html, 'html.parser')
        style_parser = StyleParser()

        for element in soup.children:
            if isinstance(element, Tag):
                self._parse_element(element, style_parser)

    def _parse_element(self, element: Tag, style_parser: StyleParser):
        """Parse a single HTML element."""
        if element.name is None:
            text = str(element).strip()
            if text:
                self.elements.append(DocumentElement(
                    type=ElementType.PARAGRAPH,
                    content=text,
                    html=text
                ))
            return

        # Get styles
        styles = {}
        if element.get('style'):
            styles = style_parser.parse_style_string(element.get('style'))

        # Page setup
        if element.name == 'page-setup':
            attrs = {k: v for k, v in element.attrs.items()}
            self.page_setup = style_parser.parse_page_setup(attrs)
            return

        # Header/Footer
        if element.name == 'header':
            self.header_content = element.get_text()
            return
        if element.name == 'footer':
            self.footer_content = element.get_text()
            return

        # Headings
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            self.elements.append(DocumentElement(
                type=ElementType.HEADING,
                content=element.get_text(),
                html=str(element),
                level=level,
                styles=styles
            ))

        # Paragraphs
        elif element.name == 'p':
            self.elements.append(DocumentElement(
                type=ElementType.PARAGRAPH,
                content=element.get_text(),
                html=str(element),
                styles=styles
            ))

        # Lists
        elif element.name in ['ul', 'ol']:
            list_type = 'bullet' if element.name == 'ul' else 'number'
            for li in element.find_all('li', recursive=False):
                li_styles = {}
                if li.get('style'):
                    li_styles = style_parser.parse_style_string(li.get('style'))
                self.elements.append(DocumentElement(
                    type=ElementType.LIST_ITEM,
                    content=li.get_text(),
                    html=str(li),
                    list_type=list_type,
                    styles=li_styles
                ))

        # Blockquote
        elif element.name == 'blockquote':
            self.elements.append(DocumentElement(
                type=ElementType.BLOCKQUOTE,
                content=element.get_text(),
                html=str(element),
                styles=styles
            ))

        # Table
        elif element.name == 'table':
            table_elem = DocumentElement(
                type=ElementType.TABLE,
                content='',
                html=str(element),
                styles=styles
            )
            # Parse table structure
            for row in element.find_all('tr'):
                row_children = []
                for cell in row.find_all(['td', 'th']):
                    row_children.append(DocumentElement(
                        type=ElementType.PARAGRAPH,
                        content=cell.get_text(),
                        html=str(cell)
                    ))
                if row_children:
                    table_elem.children.append(DocumentElement(
                        type=ElementType.PARAGRAPH,
                        content='|'.join(c.content for c in row_children),
                        html=str(row),
                        children=row_children
                    ))
            self.elements.append(table_elem)

        # Container elements
        elif element.name in ['div', 'section', 'article', 'span']:
            for child in element.children:
                if isinstance(child, Tag):
                    self._parse_element(child, style_parser)

    def get_hash(self) -> str:
        """Get hash of entire document structure."""
        content = ''.join(e.get_hash() for e in self.elements)
        return hashlib.md5(content.encode()).hexdigest()

    def get_element_hashes(self) -> List[str]:
        """Get list of element hashes."""
        return [e.get_hash() for e in self.elements]


class DocumentDiff:
    """Calculate differences between two document structures."""

    def diff(self, old: DocumentStructure, new: DocumentStructure) -> List[Change]:
        """Calculate changes between old and new document."""
        changes = []

        old_hashes = old.get_element_hashes()
        new_hashes = new.get_element_hashes()

        # Use LCS (Longest Common Subsequence) for diff
        lcs = self._lcs(old_hashes, new_hashes)

        old_idx = 0
        new_idx = 0
        lcs_idx = 0

        while old_idx < len(old_hashes) or new_idx < len(new_hashes):
            if lcs_idx < len(lcs):
                # Skip to next LCS element
                while old_idx < len(old_hashes) and old_hashes[old_idx] != lcs[lcs_idx]:
                    changes.append(Change(
                        type=ChangeType.DELETE,
                        index=old_idx,
                        old_element=old.elements[old_idx]
                    ))
                    old_idx += 1

                while new_idx < len(new_hashes) and new_hashes[new_idx] != lcs[lcs_idx]:
                    changes.append(Change(
                        type=ChangeType.INSERT,
                        index=new_idx,
                        new_element=new.elements[new_idx]
                    ))
                    new_idx += 1

                # Skip the matching element
                if old_idx < len(old_hashes) and new_idx < len(new_hashes):
                    # Check if styles changed
                    if old.elements[old_idx].styles != new.elements[new_idx].styles:
                        changes.append(Change(
                            type=ChangeType.MODIFY,
                            index=new_idx,
                            old_element=old.elements[old_idx],
                            new_element=new.elements[new_idx]
                        ))
                    old_idx += 1
                    new_idx += 1
                    lcs_idx += 1
            else:
                # No more LCS elements
                while old_idx < len(old_hashes):
                    changes.append(Change(
                        type=ChangeType.DELETE,
                        index=old_idx,
                        old_element=old.elements[old_idx]
                    ))
                    old_idx += 1

                while new_idx < len(new_hashes):
                    changes.append(Change(
                        type=ChangeType.INSERT,
                        index=new_idx,
                        new_element=new.elements[new_idx]
                    ))
                    new_idx += 1

        return changes

    def _lcs(self, a: List[str], b: List[str]) -> List[str]:
        """Find longest common subsequence."""
        m, n = len(a), len(b)
        dp = [[0] * (n + 1) for _ in range(m + 1)]

        for i in range(1, m + 1):
            for j in range(1, n + 1):
                if a[i-1] == b[j-1]:
                    dp[i][j] = dp[i-1][j-1] + 1
                else:
                    dp[i][j] = max(dp[i-1][j], dp[i][j-1])

        # Backtrack to find LCS
        result = []
        i, j = m, n
        while i > 0 and j > 0:
            if a[i-1] == b[j-1]:
                result.append(a[i-1])
                i -= 1
                j -= 1
            elif dp[i-1][j] > dp[i][j-1]:
                i -= 1
            else:
                j -= 1

        return result[::-1]


class IncrementalConverter:
    """Convert HTML to docx with incremental update support."""

    def __init__(self):
        self.style_parser = StyleParser()
        self._cache: Dict[str, Tuple[str, DocumentStructure]] = {}  # doc_id -> (docx_path, structure)

    def convert(self, html: str, doc_id: str = None) -> str:
        """
        Convert HTML to docx, using incremental update if possible.

        Args:
            html: HTML content
            doc_id: Document ID for caching

        Returns:
            Path to generated docx file
        """
        new_structure = DocumentStructure(html)

        # Check cache for incremental update
        if doc_id and doc_id in self._cache:
            cached_path, old_structure = self._cache[doc_id]

            # Check if cached file exists
            if os.path.exists(cached_path):
                diff = DocumentDiff()
                changes = diff.diff(old_structure, new_structure)

                # If changes are small, do incremental update
                change_ratio = len(changes) / max(len(new_structure.elements), 1)
                if change_ratio < 0.5:  # Less than 50% changed
                    try:
                        output_path = self._apply_incremental(cached_path, changes, new_structure)
                        self._cache[doc_id] = (output_path, new_structure)
                        return output_path
                    except Exception:
                        pass  # Fall back to full conversion

        # Full conversion
        output_path = self._convert_full(new_structure)

        if doc_id:
            self._cache[doc_id] = (output_path, new_structure)

        return output_path

    def _convert_full(self, structure: DocumentStructure) -> str:
        """Perform full HTML to docx conversion."""
        doc = Document()

        # Apply page setup
        if structure.page_setup:
            section = doc.sections[0]
            self.style_parser.apply_page_setup(section, structure.page_setup)

        # Add header
        if structure.header_content:
            section = doc.sections[0]
            header = section.header
            header.paragraphs[0].text = structure.header_content

        # Add footer
        if structure.footer_content:
            section = doc.sections[0]
            footer = section.footer
            footer.paragraphs[0].text = structure.footer_content

        # Add elements
        for element in structure.elements:
            self._add_element(doc, element)

        # Save
        temp_dir = tempfile.gettempdir()
        output_path = os.path.join(temp_dir, f'{uuid.uuid4()}.docx')
        doc.save(output_path)

        return output_path

    def _apply_incremental(self, base_path: str, changes: List[Change],
                          new_structure: DocumentStructure) -> str:
        """Apply incremental changes to existing docx."""
        doc = Document(base_path)

        # For simplicity, we'll rebuild the document body
        # but preserve styles and page setup

        # Clear existing content
        for para in list(doc.paragraphs):
            p = para._element
            p.getparent().remove(p)

        for table in list(doc.tables):
            t = table._element
            t.getparent().remove(t)

        # Add new elements
        for element in new_structure.elements:
            self._add_element(doc, element)

        # Save to new file
        temp_dir = tempfile.gettempdir()
        output_path = os.path.join(temp_dir, f'{uuid.uuid4()}.docx')
        doc.save(output_path)

        return output_path

    def _add_element(self, doc: Document, element: DocumentElement):
        """Add a document element to the docx."""
        if element.type == ElementType.HEADING:
            heading = doc.add_heading(element.content, level=element.level)
            if element.styles:
                para_style = self.style_parser.parse_paragraph_style(element.styles)
                self.style_parser.apply_paragraph_style(heading, para_style)

        elif element.type == ElementType.PARAGRAPH:
            para = doc.add_paragraph()
            self._add_inline_content(para, element)
            if element.styles:
                para_style = self.style_parser.parse_paragraph_style(element.styles)
                self.style_parser.apply_paragraph_style(para, para_style)

        elif element.type == ElementType.LIST_ITEM:
            style = 'List Bullet' if element.list_type == 'bullet' else 'List Number'
            para = doc.add_paragraph(element.content, style=style)
            if element.styles:
                para_style = self.style_parser.parse_paragraph_style(element.styles)
                self.style_parser.apply_paragraph_style(para, para_style)

        elif element.type == ElementType.BLOCKQUOTE:
            para = doc.add_paragraph(element.content)
            para.style = 'Quote'

        elif element.type == ElementType.TABLE:
            self._add_table(doc, element)

    def _add_inline_content(self, paragraph, element: DocumentElement):
        """Add inline content with styles to paragraph."""
        soup = BeautifulSoup(element.html, 'html.parser')

        # Find the main element (p, div, etc.)
        main_elem = soup.find(['p', 'div', 'span'])
        if main_elem:
            self._process_inline_children(paragraph, main_elem)
        else:
            paragraph.add_run(element.content)

    def _process_inline_children(self, paragraph, element):
        """Process inline children of an element."""
        for child in element.children:
            if child.name is None:
                # Text node
                text = str(child)
                if text:
                    paragraph.add_run(text)
            elif child.name in ['strong', 'b']:
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name in ['em', 'i']:
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif child.name == 'u':
                run = paragraph.add_run(child.get_text())
                run.underline = True
            elif child.name == 'span':
                run = paragraph.add_run(child.get_text())
                if child.get('style'):
                    styles = self.style_parser.parse_style_string(child.get('style'))
                    font_style = self.style_parser.parse_font_style(styles)
                    self.style_parser.apply_font_style(run, font_style)
            elif child.name == 'a':
                paragraph.add_run(child.get_text())
            else:
                paragraph.add_run(child.get_text())

    def _add_table(self, doc: Document, element: DocumentElement):
        """Add a table to the document."""
        if not element.children:
            return

        rows = len(element.children)
        cols = max(len(row.children) for row in element.children) if element.children else 0

        if rows == 0 or cols == 0:
            return

        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'

        for i, row_elem in enumerate(element.children):
            for j, cell_elem in enumerate(row_elem.children):
                if j < cols:
                    table.rows[i].cells[j].text = cell_elem.content

    def clear_cache(self, doc_id: str = None):
        """Clear cache for a document or all documents."""
        if doc_id:
            if doc_id in self._cache:
                del self._cache[doc_id]
        else:
            self._cache.clear()

    def get_cache_info(self, doc_id: str) -> Optional[Dict[str, Any]]:
        """Get cache information for a document."""
        if doc_id in self._cache:
            path, structure = self._cache[doc_id]
            return {
                'cached_path': path,
                'element_count': len(structure.elements),
                'hash': structure.get_hash()
            }
        return None


# Global converter instance
_converter_instance: Optional[IncrementalConverter] = None


def get_converter() -> IncrementalConverter:
    """Get global converter instance."""
    global _converter_instance
    if _converter_instance is None:
        _converter_instance = IncrementalConverter()
    return _converter_instance
